import docx
import re

# ===================== Snow ===========================
snow_2 = [" ", "\t"]


def snow_encode(msg, character_set=None, binary=False):
    code = ''
    if not character_set:
        character_set = snow_2
    if binary:
        msg_bytes = msg
    else:
        msg_bytes = bytes(msg, 'utf-8')
    for by in msg_bytes:
        bit_mask = 0b00000001
        for i in range(8):
            m_byte = by & bit_mask
            by = by >> 1
            if m_byte == 1:
                code += character_set[1]
            elif m_byte == 0:
                code += character_set[0]

    return code


def snow_decode(code, character_set=None, binary=False):
    msg_bytes = []
    if not character_set:
        character_set = snow_2

    # range behave a little weird here is there a less hacky solution?
    for i in range(8, len(code) + 1, 8):
        m_byte = 0b00000000
        encoded_byte = code[i - 8:i]
        # for char in encoded_byte:
        #     print("char: {} ord: {}".format(char, hex(ord(char))))

        for j in range(8):
            m_byte = m_byte >> 1
            if encoded_byte[j] == character_set[1]:
                m_byte = m_byte | 0b10000000
            elif encoded_byte[j] == character_set[0]:
                pass

        msg_bytes.append(m_byte)

    if binary:
        return bytes(msg_bytes)
    return bytes(msg_bytes).decode('utf-8')


def encode(unencoded_string, msg, method="snow", binary=False, replacements=None, delimiter=None):
    '''
    Main encoding method
    Dispatches to corresponding encoder based on specified method and handles
    insertion/appending etc. of message into the string.
    '''

    if method == "snow":
        if not delimiter:
            delimiter = '\t\t\t'
        code = snow_encode(msg, character_set=replacements, binary=binary)
        return unencoded_string + delimiter + code

    else:
        raise Exception("Method: {}, is not supported".format(method))


def decode(encoded_string, method="snow", binary=False, replacements=None, delimiter=None):
    '''
    Main decoding method
    Dispatches to corresponding decoder based on specified method and handles
    extraction of encoded message from the string.
    '''
    if method == "snow":
        if not delimiter:
            delimiter = '\t\t\t'
        regex = "{}(.+)$".format(delimiter)
        m = re.search(regex, encoded_string)
        code = m.groups()[0]

        return snow_decode(code, character_set=replacements, binary=binary)


# ENCODING
# Step1 (ENCODING)
def check_documentFileType(file_path):
    documentTypeChecker = ""
    documentFile_types = file_path

    # Document(Words Fromat) File
    if documentFile_types.endswith('.docx'):
        words_documentFile = docx.Document(documentFile_types)

        # Add Document (Words Format) Type
        documentTypeChecker += "docx"

        text = ''
        # Convert file (Words Fromat) into text variable   
        for para in words_documentFile.paragraphs:
            text += para.text + '\n'

    # Document(Text Format) File
    elif documentFile_types.endswith('.txt'):

        # Add Document (Text Format) Type
        documentTypeChecker += "txt"

        # Opening the Text file
        with open(documentFile_types, 'r') as text_documentFile:

            # Convert file (Text Format) into text variable   
            text = text_documentFile.read()

    # Return extracted content and type from document file
    return text, documentTypeChecker


# Step2 (ENCODING)
def open_secretFile(file_path):
    secretMessage = file_path

    # opening the text file
    with open(secretMessage, 'r') as text_secretMessage:
        # Reading each line    
        secret_msg = text_secretMessage.read()

    return secret_msg


def savedEncodedFile(encoded, documentTypeChecker, filepath):
    # Save Words Type
    if documentTypeChecker == "docx":
        document = docx.Document()
        document.add_paragraph(encoded)
        document.save(f'{filepath}/EncodedWordsFile.' + documentTypeChecker)

        print("Encoded Words File Created")

    # Save Text Type
    elif documentTypeChecker == "txt":
        with open(f'{filepath}/encodedTextFile.' + documentTypeChecker, 'w') as file:
            file.write(encoded)
        print("Encoded Text File Created")


# Step3 (ENCODING)
def encode_text_to_text_stega(text, secret_msg, documentTypeChecker, filepath):
    encoded = encode(text, secret_msg)

    # Save document file encoded with secret message
    savedEncodedFile(encoded, documentTypeChecker, filepath)


# DECODING
# Step1 (DECODING)
def check_EncodedDocumentFile(file_path):
    # Getting the input of Document file (eg. Word, .txt, .xls)
    encodedDocumentFile_types = file_path

    if encodedDocumentFile_types.endswith('.docx'):
        encodedWords_documentFile = docx.Document(encodedDocumentFile_types)
        encodedText = ''

        # Convert file (Words Fromat) into text variable   
        for para in encodedWords_documentFile.paragraphs:
            encodedText += para.text + '\n'

    # Document(Text Format) File
    elif encodedDocumentFile_types.endswith('.txt'):

        # opening the Text file
        with open(encodedDocumentFile_types, 'r') as encodedText_documentFile:
            # Convert file (Text Format) into text variable   
            encodedText = encodedText_documentFile.read()

    return encodedText


def decodedSecretMessage(encoded, filepath):
    with open(f'{filepath}/decodedSecretText.txt', 'w') as file:
        file.write(encoded)
    print("Decoded secret message from Encoded Document File")


# Step2 (DECODING)
def decode_text_to_text_stega(encodedText, filepath):
    decoded = decode(encodedText)
    decodedSecretMessage(decoded, filepath)


# These 2 are the only function u need to call
# Text to Text encoding function
def t2t_encoding_steganoFunction(file_path_document_file, file_path_secret_file, output_filepath):
    extracted_text, document_type = check_documentFileType(file_path_document_file)
    extracted_secret_text = open_secretFile(file_path_secret_file)
    encode_text_to_text_stega(extracted_text, extracted_secret_text, document_type, output_filepath)
    return 'File successfully encoded!'


# Text to Text decoding function
def t2t_decoding_steganoFunction(file_path_encoded_document_file, filepath):
    extracted_encoded_text = check_EncodedDocumentFile(file_path_encoded_document_file)
    decode_text_to_text_stega(extracted_encoded_text, filepath)
    return 'File successfully decoded!'
